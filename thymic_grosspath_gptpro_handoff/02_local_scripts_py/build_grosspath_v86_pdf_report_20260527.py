from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\影响分析")
V86_DIR = ROOT / "outputs" / "grosspath_rc_v86_paper_ready_summary_pack_20260527"
OUT_DIR = ROOT / "汇报"
OUT_DOCX = OUT_DIR / "2026-05-27_Task7风险可控跨域诊断框架阶段汇报_v50-v86.docx"


BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
DARK = "222222"
MUTED = "666666"


WORKFLOW_ZH = {
    "Standard risk-control workflow (v50)": "v50 标准风险控制",
    "High-risk miss protection (v75)": "v75 高危漏诊保护",
    "Bidirectional light guard (v79-light)": "v79-light 双向轻量保护",
    "Bidirectional strict guard (v79-strict)": "v79-strict 双向严格保护",
}

DOMAIN_ZH = {
    "old_data": "旧数据",
    "third_batch": "第三批",
    "strict_external": "严格外部集",
}

CLAIM_ZH = {
    "main": "主结果",
    "candidate": "候选",
}


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_run(paragraph, text: str, bold: bool = False, size: float | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_table(table, header_fill: str = LIGHT_GRAY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_east_asia_font(r)
                    r.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor.from_string(DARK)


def put_cell(cell, text: str, bold: bool = False, align: str = "center", size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Task7 风险可控跨域诊断框架阶段汇报")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    add_run(p, "v50-v86 结果汇总 | 2026-05-27 | 胸腺大体病理图像二分类", size=10, color=MUTED)

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(
        p,
        "本阶段我们的重点已经从单纯追求一个分类器分数，转为构建“自动诊断 + 风险门控 + 复核触发”的可部署流程。"
        "目前最稳的主推版本是 v79-light：在严格外部集上 BAcc 99.18%，高危漏诊 0 例，低危误升级 1 例，控制率 82.41%。"
        "v79-strict 在该外部集上达到 0 错例，但仍按候选高安全版本处理，后续需要更多外部批次验证。",
        bold=True,
        size=10.5,
        color=DARK,
    )
    style_table(box, header_fill=LIGHT_BLUE)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "• ", bold=True, color=BLUE)
        add_run(p, item)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        put_cell(table.cell(0, i), h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = "left" if i == 0 else "center"
            put_cell(cells[i], value, align=align)
    set_table_width(table, widths_cm)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def make_three_domain_rows(primary: pd.DataFrame) -> list[list[str]]:
    keep = primary[primary["policy_id"].isin(["v50_main", "v75_quality_lowconf", "v79_light_lowrisk_guard", "v79_strict_lowrisk_guard"])]
    rows = []
    for _, r in keep.iterrows():
        rows.append(
            [
                DOMAIN_ZH.get(r["domain"], r["domain"]),
                WORKFLOW_ZH.get(r["workflow"], r["workflow"]),
                r["control_rate"],
                r["balanced_accuracy"],
                r["sensitivity"],
                r["specificity"],
                str(int(r["fn"])),
                str(int(r["fp"])),
            ]
        )
    return rows


def make_strict_rows(strict: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, r in strict.iterrows():
        rows.append(
            [
                WORKFLOW_ZH.get(r["workflow"], r["workflow"]),
                r["control_rate"],
                f"{r['balanced_accuracy']}\n{r['balanced_accuracy_wilson_95ci']}",
                f"{r['sensitivity']}\n{r['sensitivity_wilson_95ci']}",
                f"{r['specificity']}\n{r['specificity_wilson_95ci']}",
                str(int(r["fn"])),
                str(int(r["fp"])),
                CLAIM_ZH.get(r["claim_tier"], r["claim_tier"]),
            ]
        )
    return rows


def make_ablation_rows(ablation: pd.DataFrame) -> list[list[str]]:
    keep = ablation[ablation["domain"].isin(["third_batch", "strict_external"])]
    rows = []
    for _, r in keep.iterrows():
        rows.append(
            [
                DOMAIN_ZH.get(r["domain"], r["domain"]),
                r["module_label"],
                r["delta_control_rate"],
                r["delta_bacc"],
                str(int(r["delta_fn"])),
                str(int(r["delta_fp"])),
            ]
        )
    return rows


def make_paired_rows(paired: pd.DataFrame) -> list[list[str]]:
    keep = paired[
        paired["comparison"].isin(
            [
                "Light full workflow vs v50",
                "Strict full workflow vs v50",
                "Light low-risk guard vs v75",
            ]
        )
    ]
    rows = []
    for _, r in keep.iterrows():
        rows.append(
            [
                DOMAIN_ZH.get(r["domain"], r["domain"]),
                r["comparison"],
                r["delta_bacc"],
                r["delta_bacc_bootstrap_95ci"],
                str(int(r["delta_fn"])),
                str(int(r["delta_fp"])),
                str(r["mcnemar_p"]),
            ]
        )
    return rows


def build_report() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    primary = pd.read_csv(V86_DIR / "v86_primary_fixed_workflow_table.csv")
    strict = pd.read_csv(V86_DIR / "v86_strict_external_focus_table.csv")
    adaptive = pd.read_csv(V86_DIR / "v86_adaptive_workflow_table.csv")
    ablation = pd.read_csv(V86_DIR / "v86_module_ablation_table.csv")
    paired = pd.read_csv(V86_DIR / "v86_paired_delta_stats_table.csv")

    doc = Document()
    configure_doc(doc)
    add_title_block(doc)

    doc.add_heading("一、目前最能对外讲清楚的结论", level=1)
    add_bullets(
        doc,
        [
            "我们不把当前结果包装成单一模型的偶然高分，而是把它定义为风险可控的大体病理图像诊断流程：先自动判断，再根据低置信、质量偏移、误升级风险触发复核。",
            "v50 是稳定基础流程，在旧数据、第三批、严格外部集上分别达到 99.65%、95.29%、97.30% BAcc，说明基础二阶段框架已经能跨批次工作。",
            "v75 的主要价值是减少高危漏诊；v79 的主要价值是进一步减少低危误升级。两个模块解决的是不同错误类型，不是简单叠模型。",
            "v79-light 是当前主推高安全版本；v79-strict 结果更好，但因为严格外部集样本量有限，现阶段只作为高安全候选版本。",
        ],
    )

    doc.add_heading("二、三域固定工作流结果", level=1)
    doc.add_paragraph(
        "下表保留旧数据、第三批和严格外部集三组结果。我们更看重跨域一致性，而不是单一数据集的最高点估计。"
    )
    add_simple_table(
        doc,
        ["数据域", "流程", "控制率", "BAcc", "Sens", "Spec", "FN", "FP"],
        make_three_domain_rows(primary),
        [1.5, 4.2, 1.5, 1.5, 1.5, 1.5, 0.8, 0.8],
    )

    doc.add_heading("三、严格外部集焦点结果", level=1)
    doc.add_paragraph(
        "严格外部集是最关键的泛化检验。这里我们同时报告点估计和 Wilson 95% 置信区间，避免把小样本 100% 误读成真实世界确定性。"
    )
    add_simple_table(
        doc,
        ["流程", "控制率", "BAcc\n95%CI", "Sens\n95%CI", "Spec\n95%CI", "FN", "FP", "定位"],
        make_strict_rows(strict),
        [3.3, 1.35, 2.0, 2.0, 2.0, 0.65, 0.65, 1.0],
    )

    doc.add_heading("四、模块贡献拆解", level=1)
    doc.add_paragraph(
        "我们把高危漏诊保护和低危误升级保护拆开评估，目的是真正说明每个模块承担什么临床风险。"
    )
    add_simple_table(
        doc,
        ["数据域", "模块", "控制率变化", "BAcc变化", "FN变化", "FP变化"],
        make_ablation_rows(ablation),
        [1.6, 4.1, 1.5, 1.5, 1.0, 1.0],
    )

    doc.add_heading("五、配对统计和表述边界", level=1)
    doc.add_paragraph(
        "同病例配对比较更接近真实增益判断。第三批上 v79-light 相比 v50 的 BAcc 提升有正向区间支持；严格外部集错误数明显减少，但受样本量限制，统计上应写成趋势和错误减少，而不是强显著。"
    )
    add_simple_table(
        doc,
        ["数据域", "比较", "BAcc变化", "bootstrap 95%CI", "FN变化", "FP变化", "McNemar p"],
        make_paired_rows(paired),
        [1.45, 3.2, 1.35, 2.3, 1.0, 1.0, 1.2],
    )

    doc.add_heading("六、部署型流程和论文贡献", level=1)
    adaptive_light = adaptive[adaptive["policy_id"] == "adaptive_light_on_shift"].iloc[0]
    fixed_light = adaptive[adaptive["policy_id"] == "v79_light_lowrisk_guard"].iloc[0]
    doc.add_paragraph(
        f"如果只追求已有三域的最高安全性，固定 v79-light/strict 更强；如果考虑真实部署，v82 的无标签批次审计更有方法学价值。"
        f"例如 adaptive v50->v79-light 的全域控制率为 {adaptive_light['overall_control_rate']}，低于固定 v79-light 的 {fixed_light['overall_control_rate']}，"
        f"但能在识别到严重批次偏移时自动切换安全流程。"
    )
    add_bullets(
        doc,
        [
            "计算机侧主线：无标签批次审计识别采集域偏移，再按风险等级选择不同复核强度。",
            "医学侧主线：高危漏诊和低危误升级分开控制，结果可以对应临床安全诉求。",
            "论文写法：v50 是稳定基线，v79-light 是主推高安全流程，v79-strict 是候选上限，v82 是部署式自适应框架。",
        ],
    )

    doc.add_heading("七、不能夸大的部分", level=1)
    add_bullets(
        doc,
        [
            "严格外部集已经被我们用于大量分析，后续论文中必须区分正式盲法验证结果和暴露后的探索性发现。",
            "v79-strict 的 0 错例是很有价值的候选结果，但 Wilson 区间下界仍约 93%，不能直接写成真实世界一定 100%。",
            "质量门控和外部域偏移解释是合理方向，但还需要更多外部医院批次来验证阈值稳定性。",
        ],
    )

    doc.add_heading("八、下一步", level=1)
    add_bullets(
        doc,
        [
            "继续把质量、视图、批次偏移做成开发集可标定的模块，减少依赖外部集分数反向调参。",
            "补充更多外部批次或前瞻性样本，验证 v79-light/v79-strict 的安全性是否稳定。",
            "把 v82-v85 的统计、消融和决策效用图表整理成论文 Results 的固定主表和补充材料。",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(footer, "Task7 风险可控跨域诊断框架阶段汇报 | v50-v86", size=8, color=MUTED)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    path = build_report()
    print(path)
