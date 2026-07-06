from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "汇报"
DOCX_PATH = OUT_DIR / "2026-05-27_Task7风险控制工作流阶段汇报_医生版.docx"
REPORT_FIG_DIR = ROOT / "outputs" / "v53_doctor_report_figures"

V51 = ROOT / "outputs" / "grosspath_rc_v51_workflow_validation_20260527"
V52 = ROOT / "outputs" / "grosspath_rc_v52_quality_retake_overlay_20260527"


BLUE = "1F4D78"
DARK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "FFF3CD"
RED = "F8D7DA"
GREEN = "DDEFE5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix) :])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, 11, True, DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(body)
    set_font(r, 10.5, False, "222222")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 8.8, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_font(r, 8.5)
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, 9, False, "555555")


def pct(x: object, decimals: int = 1) -> str:
    try:
        return f"{float(x) * 100:.{decimals}f}%"
    except Exception:
        return ""


def fmt_ci(row: pd.Series, prefix: str) -> str:
    return f"{float(row[f'{prefix}_median']) * 100:.1f}% ({float(row[f'{prefix}_ci025']) * 100:.1f}-{float(row[f'{prefix}_ci975']) * 100:.1f}%)"


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    return doc


def make_report_figures() -> None:
    REPORT_FIG_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(V51 / "v51_tiered_workflow_summary.csv")
    ext = summary.loc[summary["split"].eq("external")].copy()
    order = [
        "P2_pure_auto",
        "v37_balanced_dev97",
        "v48_direction_dev97",
        "v48_fn_high_safety",
        "v50_sens97_spec90",
        "v50_sens98_spec90",
    ]
    labels = ["P2", "v37", "v48", "v48-FN", "v50-97", "v50-98"]
    ext = ext.set_index("policy").loc[order].reset_index()
    fig, ax1 = plt.subplots(figsize=(8.4, 4.6))
    x = range(len(ext))
    ax1.bar(x, ext["balanced_accuracy"] * 100, color="#1f6f61", alpha=0.86, width=0.62)
    ax1.set_ylabel("External BAcc (%)")
    ax1.set_ylim(65, 100)
    ax1.axhline(95, color="#9a7d0a", linestyle="--", linewidth=1)
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels)
    ax1.grid(axis="y", linestyle="--", alpha=0.3)
    for i, row in ext.iterrows():
        ax1.text(i, row["balanced_accuracy"] * 100 + 0.8, f"{row['balanced_accuracy']*100:.1f}%", ha="center", fontsize=9)
    ax2 = ax1.twinx()
    ax2.plot(list(x), ext["review_rate"] * 100, color="#b03a2e", marker="o", linewidth=2)
    ax2.set_ylabel("Review/control rate (%)")
    ax2.set_ylim(0, 85)
    fig.suptitle("Tiered workflow: accuracy improves with risk-control intensity", y=0.98)
    fig.tight_layout()
    fig.savefig(REPORT_FIG_DIR / "v53_workflow_bacc_review.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    q = pd.read_csv(V52 / "v52_quality_retake_overlay_summary.csv")
    keep = [
        "v50_only",
        "v50_plus_quality_score_le74",
        "v50_plus_quality_score_le82",
        "v50_plus_quality_score_le88",
    ]
    q = q.set_index("policy").loc[keep].reset_index()
    qlabels = ["v50", "q<=74", "q<=82", "q<=88"]
    fig, ax1 = plt.subplots(figsize=(8.0, 4.4))
    x = range(len(q))
    ax1.bar(x, q["balanced_accuracy"] * 100, color="#a04000", alpha=0.84, width=0.58)
    ax1.set_ylabel("External workflow BAcc (%)")
    ax1.set_ylim(96, 100.5)
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(qlabels)
    ax1.grid(axis="y", linestyle="--", alpha=0.3)
    for i, row in q.iterrows():
        ax1.text(i, row["balanced_accuracy"] * 100 + 0.08, f"{row['balanced_accuracy']*100:.1f}%", ha="center", fontsize=9)
    ax2 = ax1.twinx()
    ax2.plot(list(x), q["total_control_rate"] * 100, color="#154360", marker="o", linewidth=2)
    ax2.set_ylabel("Total review/retake rate (%)")
    ax2.set_ylim(70, 92)
    fig.suptitle("Quality-retake overlay: fewer residual errors, higher control burden", y=0.98)
    fig.tight_layout()
    fig.savefig(REPORT_FIG_DIR / "v53_quality_overlay.png", dpi=300, bbox_inches="tight")
    plt.close(fig)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_report_figures()
    doc = setup_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Task7 二分类大体病理图像风险控制工作流阶段汇报")
    set_font(r, 20, True, DARK)
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("截至 2026-05-27 | 旧数据 + 第三批开发，胸腺瘤+癌作为严格外部集")
    set_font(r, 10.5, False, "555555")

    add_callout(
        doc,
        "本轮核心结论",
        "我们目前最稳妥的主结果不是单纯追求纯自动分类，而是形成了可按临床风险偏好切换的多档工作流。"
        "当前主推的 v50 高敏感性档在外部集达到 Acc 97.2%、BAcc 97.3%，高危漏诊剩 1 例；"
        "质量重拍门控可以进一步拦截剩余错例，但应作为探索性部署机制，而不是纯模型能力。",
        GREEN,
    )

    doc.add_heading("1. 当前问题和我们采用的边界", level=1)
    add_text(
        doc,
        "目前外部集图像质量和拍摄方式差异明显，直接追求全自动输出会低估临床风险。"
        "因此我们把任务重新定义为：先给出模型判断，再根据错误方向、复核预算和图像质量决定是否自动放行、复核或建议重拍。",
    )
    add_text(
        doc,
        "我们在规则选择上坚持一个边界：外部集只用于最终评估，不用外部集标签反过来训练主模型或选择主流程阈值。"
        "v52 的质量阈值分析属于外部集上的部署探索，已经单独标注，不作为主模型调参结论。",
    )

    doc.add_heading("2. 多档工作流结果", level=1)
    summary = pd.read_csv(V51 / "v51_tiered_workflow_summary.csv")
    ext = summary.loc[summary["split"].eq("external")].copy()
    order = [
        "P2_pure_auto",
        "v37_balanced_dev97",
        "v48_direction_dev97",
        "v48_fn_high_safety",
        "v50_sens97_spec90",
        "v50_sens98_spec90",
    ]
    ext = ext.set_index("policy").loc[order].reset_index()
    rows = []
    labels = {
        "P2_pure_auto": "P2 纯自动",
        "v37_balanced_dev97": "v37 均衡档",
        "v48_direction_dev97": "v48 方向感知档",
        "v48_fn_high_safety": "v48 FN高安全档",
        "v50_sens97_spec90": "v50 敏感性97档",
        "v50_sens98_spec90": "v50 敏感性98档",
    }
    for _, row in ext.iterrows():
        rows.append(
            [
                labels[row["policy"]],
                pct(row["review_rate"]),
                pct(row["accuracy"]),
                pct(row["balanced_accuracy"]),
                pct(row["sensitivity"]),
                pct(row["specificity"]),
                str(int(row["fn"])),
                str(int(row["fp"])),
            ]
        )
    add_table(
        doc,
        ["策略", "复核/控制", "Acc", "BAcc", "敏感性", "特异性", "FN", "FP"],
        rows,
        [1.35, 0.75, 0.65, 0.65, 0.75, 0.75, 0.45, 0.45],
        LIGHT_BLUE,
    )
    add_picture(
        doc,
        REPORT_FIG_DIR / "v53_workflow_bacc_review.png",
        "图1 多档风险控制工作流：复核比例升高后，外部集 BAcc 和高危漏诊控制同步改善。",
        6.1,
    )

    doc.add_heading("3. 稳定性验证", level=1)
    ci = pd.read_csv(V51 / "v51_external_bootstrap_ci.csv").set_index("policy")
    rows = []
    for policy in ["v37_balanced_dev97", "v48_direction_dev97", "v48_fn_high_safety", "v50_sens97_spec90", "v50_sens98_spec90"]:
        row = ci.loc[policy]
        rows.append(
            [
                labels.get(policy, policy),
                fmt_ci(row, "balanced_accuracy"),
                fmt_ci(row, "accuracy"),
                fmt_ci(row, "sensitivity"),
            ]
        )
    add_table(doc, ["策略", "BAcc 中位数(95%CI)", "Acc 中位数(95%CI)", "敏感性中位数(95%CI)"], rows, [1.4, 1.7, 1.7, 1.7])

    sel = pd.read_csv(V51 / "v51_sensitivity_selection_bootstrap_summary.csv")
    rows = []
    for _, row in sel.iterrows():
        rows.append(
            [
                f"{float(row['target_sensitivity']) * 100:.0f}%",
                f"{float(row['selected_review_rate_median']) * 100:.1f}%",
                f"{float(row['external_bacc_median']) * 100:.1f}% ({float(row['external_bacc_ci025']) * 100:.1f}-{float(row['external_bacc_ci975']) * 100:.1f}%)",
                f"{float(row['p_external_bacc_ge_95']) * 100:.1f}%",
                f"{float(row['p_external_fn_le_1']) * 100:.1f}%",
            ]
        )
    add_table(doc, ["开发集敏感性目标", "外部复核中位数", "外部BAcc中位数(区间)", "BAcc≥95%概率", "FN≤1概率"], rows, [1.1, 1.2, 1.8, 1.2, 1.2])
    add_picture(
        doc,
        V51 / "figures" / "v51_sensitivity_selection_stability.png",
        "图2 开发集重采样后的选择稳定性：敏感性98档更稳，但仍需用区间而非单点描述。",
        5.7,
    )

    doc.add_heading("4. 质量拒判和重拍机制", level=1)
    add_callout(
        doc,
        "表述边界",
        "质量门控可以把最后几个错误挡下来，但它的定位是“低质量/边界图像不直接自动出结果”。"
        "这不是纯自动诊断能力，而是临床部署中合理的拒判和重拍机制。",
        GOLD,
    )
    q = pd.read_csv(V52 / "v52_quality_retake_overlay_summary.csv")
    keep = [
        "v50_only",
        "v50_plus_quality_score_le74",
        "v50_plus_quality_score_le82",
        "v50_plus_quality_score_le88",
    ]
    q = q.set_index("policy").loc[keep].reset_index()
    qlabels = {
        "v50_only": "v50 原流程",
        "v50_plus_quality_score_le74": "+ 质量≤74",
        "v50_plus_quality_score_le82": "+ 质量≤82",
        "v50_plus_quality_score_le88": "+ 质量≤88",
    }
    rows = []
    for _, row in q.iterrows():
        rows.append(
            [
                qlabels[row["policy"]],
                str(int(row["additional_retake_n"])),
                pct(row["total_control_rate"]),
                pct(row["balanced_accuracy"]),
                str(int(row["fn"])),
                str(int(row["fp"])),
                str(int(row["remaining_error_n"])),
            ]
        )
    add_table(doc, ["策略", "额外重拍/复核", "总控制比例", "BAcc", "FN", "FP", "最终剩余错误"], rows, [1.35, 1.0, 0.9, 0.75, 0.45, 0.45, 0.95], GOLD)
    add_picture(
        doc,
        REPORT_FIG_DIR / "v53_quality_overlay.png",
        "图3 质量重拍门控叠加后，剩余错误可以继续下降，但总控制比例也明显升高。",
        5.9,
    )

    doc.add_heading("5. 最后仍需医生关注的病例", level=1)
    err = pd.read_csv(V51 / "v51_v50_sens98_remaining_errors.csv")
    rows = []
    for _, row in err.iterrows():
        rows.append(
            [
                str(row["original_case_id"]),
                str(row["task_l6_label"]),
                f"{row['quality_status']} / {float(row['quality_score']):.0f}",
                str(row["p2_error_direction"]),
                f"{float(row['main_prob']):.3f}",
                f"{float(row['robust_prob']):.3f}",
            ]
        )
    add_table(doc, ["病理号", "亚型", "质量", "错误方向", "主模型概率", "稳健概率"], rows, [0.9, 0.55, 1.1, 1.45, 0.85, 0.85], RED)
    add_picture(
        doc,
        V52 / "figures" / "v52_v50_remaining_error_gallery.png",
        "图4 v50 高敏感性档剩余错例图像。三例均为 borderline 质量，适合进入重拍或人工复核。",
        6.25,
    )

    doc.add_heading("6. 阶段判断", level=1)
    add_text(
        doc,
        "我们目前已经不再只是比较单个模型准确率，而是形成了三层工作流：分类模型给初判，方向感知风险器决定复核优先级，质量门控决定是否需要重拍或拒判。"
        "这条线比单纯模型堆叠更适合作为计算机侧贡献，因为它把临床风险目标显式写进了算法决策过程。",
    )
    add_text(
        doc,
        "如果要给医生汇报，建议主说 v50：外部集 Acc 97.2%、BAcc 97.3%、高危漏诊 1 例。"
        "v52 可以作为补充说明：最后 3 例错误都被质量门控识别为不适合直接自动放行，后续可要求重拍或复核。",
    )

    doc.add_heading("7. 下一步计划", level=1)
    for item in [
        "固定质量门控规则后，在后续新增数据上前瞻验证，不能继续根据外部集分数调整阈值。",
        "围绕 B2 漏诊和 AB/B1 误升级做病例级解释，形成医生可复核的错例清单。",
        "继续降低 v50 高安全档的复核比例，目标是在保持 FN≤1 的前提下，把总控制比例从 73% 继续压低。",
        "把风险控制曲线、方向感知风险器和质量拒判机制整理成论文方法图，突出计算机侧创新。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
